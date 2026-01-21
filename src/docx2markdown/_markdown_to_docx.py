from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn  # QName helper for namespaces

import os
import re
from pathlib import Path

def markdown_to_docx(markdown_file, output_docx):
    """Convert a Markdown file to a .docx file."""
    doc = Document()
    
    # 获取 markdown 文件所在目录，用于解析相对路径
    md_file_dir = Path(markdown_file).parent

    with open(markdown_file, "r", encoding="utf-8") as md_file:
        lines = md_file.readlines()

    table_buffer = []  # To collect table lines
    in_table = False  # Flag for table parsing

    for line in lines:
        line = line.rstrip()

        # Table handling
        if "|" in line and "---" not in line:
            if not in_table:
                in_table = True  # Start of a table
            table_buffer.append(line)
            continue
        elif in_table and "---" in line:
            continue  # Ignore separator row
        elif in_table and line.strip() == "":
            add_table(doc, table_buffer)  # Add the parsed table to the document
            table_buffer = []  # Reset buffer
            in_table = False
            continue

        # Headings
        if line.startswith("#"):
            heading_level = len(line.split(" ", 1)[0])  # Number of # indicates heading level
            text = line.lstrip("#").strip()
            doc.add_heading(text, level=min(heading_level, 5))

        # Multi-level bullet points
        elif re.match(r"^(\s*)[-*]\s", line):
            indent_level = len(re.match(r"^(\s*)", line).group(1)) // 2
            text = line.strip("-* ").strip()
            add_bullet_point(doc, text, level=indent_level)

        # Numbered lists
        elif re.match(r"^\d+\.\s", line):
            text = line.split(". ", 1)[1]
            doc.add_paragraph(text, style="List Number")

        # Images - Markdown format: ![alt text](image_path)
        elif line.startswith("![") and "](" in line:
            alt_text = re.search(r"!\[(.*?)\]", line).group(1)
            image_path = re.search(r"\((.*?)\)", line).group(1)
            # 处理相对路径
            image_path = resolve_image_path(image_path, md_file_dir)
            
            if image_path and os.path.exists(image_path):
                try:
                    doc.add_picture(image_path, width=Inches(3.0))
                except Exception as e:
                    doc.add_paragraph(f"[Image error: {alt_text} - {str(e)}]")
            else:
                doc.add_paragraph(f"[Image not found: {alt_text}]")
        
        # Images - HTML format: <img src="..." class="icon" />
        elif "<img" in line and "src=" in line:
            # 提取图片路径
            img_match = re.search(r'<img[^>]+src=["\']([^"\']+)["\']', line)
            if img_match:
                image_path = img_match.group(1)
                # 处理相对路径
                image_path = resolve_image_path(image_path, md_file_dir)
                
                if image_path and os.path.exists(image_path):
                    try:
                        doc.add_picture(image_path, width=Inches(3.0))
                    except Exception as e:
                        doc.add_paragraph(f"[Image error: {str(e)}]")
                else:
                    doc.add_paragraph(f"[Image not found: {image_path}]")

        # Links (e.g., [text](url))
        elif "[" in line and "](" in line:
            process_line_with_links(doc, line)

        # Bold and Italic formatting
        elif "**" in line or "*" in line:
            parse_and_add_text_with_formatting(doc, line)

        # Plain text
        else:
            if line:
                doc.add_paragraph(line)

    # Save the document
    doc.save(output_docx)


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph in a Word document.

    :param paragraph: The paragraph to which the hyperlink will be added.
    :param url: The URL for the hyperlink.
    :param text: The display text for the hyperlink.
    :return: None
    """
    # Get the relationship id for the hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Create a w:r element
    run = OxmlElement("w:r")

    # Create a w:t element and set the text
    text_element = OxmlElement("w:t")
    text_element.text = text

    # Add w:t to w:r, and w:r to w:hyperlink
    run.append(text_element)
    hyperlink.append(run)

    # Add the hyperlink to the paragraph
    paragraph._element.append(hyperlink)


def process_line_with_links(doc, line):
    """
    Parse a line for hyperlinks and regular text, adding them to the document.

    :param doc: The Word document object.
    :param line: The Markdown line to parse.
    :return: None
    """
    paragraph = doc.add_paragraph()
    # Regular expression to match [text](url)
    pattern = r"\[(.*?)\]\((.*?)\)"
    cursor = 0  # To track the position in the line

    # Find all matches of the pattern
    for match in re.finditer(pattern, line):
        # Add the text before the hyperlink
        start, end = match.span()
        if cursor < start:
            paragraph.add_run(line[cursor:start])

        # Add the hyperlink
        link_text = match.group(1)
        link_url = match.group(2)
        add_hyperlink(paragraph, link_url, link_text)

        # Move the cursor past the current match
        cursor = end

    # Add any remaining text after the last hyperlink
    if cursor < len(line):
        paragraph.add_run(line[cursor:])



def add_table(doc, table_lines):
    """Add a Markdown-style table to the document."""
    # Split header and rows
    headers = table_lines[0].split("|")[1:-1]  # Extract header columns
    rows = [row.split("|")[1:-1] for row in table_lines[1:]]  # Extract data rows

    # Create a table in the document
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"

    # Add headers
    for i, header in enumerate(headers):
        table.cell(0, i).text = header.strip()

    # Add rows
    for row_idx, row in enumerate(rows):
        for col_idx, cell in enumerate(row):
            table.cell(row_idx + 1, col_idx).text = cell.strip()


def add_bullet_point(doc, text, level):
    """Add a bullet point with the appropriate indentation."""
    # Add a new paragraph with the "List Bullet" style
    paragraph = doc.add_paragraph(text, style="List Bullet")

    # Access the XML of the paragraph
    p = paragraph._element

    # Create or find the <w:numPr> element
    numPr = p.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        p.insert(0, numPr)

    # Set the numbering level <w:ilvl>
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        numPr.append(ilvl)
    ilvl.set(qn("w:val"), str(level))  # Set the level (e.g., 0, 1, 2)

    # Set the numbering ID <w:numId> (required for proper bullet point rendering)
    numId = numPr.find(qn("w:numId"))
    if numId is None:
        numId = OxmlElement("w:numId")
        numPr.append(numId)
    numId.set(qn("w:val"), "1")  # Use numbering ID 1 (default for List Bullet)


def resolve_image_path(image_path, md_file_dir):
    """
    解析图片路径，支持相对路径和绝对路径。
    
    :param image_path: 图片路径（可能是相对路径或绝对路径）
    :param md_file_dir: Markdown 文件所在目录
    :return: 解析后的绝对路径，如果路径无效返回 None
    """
    if not image_path:
        return None
    
    # 移除开头的 ./
    if image_path.startswith("./"):
        image_path = image_path[2:]
    
    # 如果是绝对路径，直接返回
    if os.path.isabs(image_path):
        return image_path if os.path.exists(image_path) else None
    
    # 相对路径：相对于 markdown 文件所在目录
    full_path = md_file_dir / image_path
    if full_path.exists():
        return str(full_path)
    
    return None


def parse_and_add_text_with_formatting(doc, line):
    """
    Parse a line for bold (**text**) and italic (*text*) formatting
    and add it to the document.
    """
    paragraph = doc.add_paragraph()

    # Define a regex pattern to match bold and italic text
    pattern = r"(\*\*.*?\*\*|\*.*?\*)"

    # Split the line by bold/italic patterns while keeping the delimiters
    parts = re.split(pattern, line)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Bold text
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            # Italic text
            paragraph.add_run(part[1:-1]).italic = True
        else:
            # Regular text
            paragraph.add_run(part)




