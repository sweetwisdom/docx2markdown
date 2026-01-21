import docx
import os
import re
import uuid
import time
import random
from lxml import etree
from pathlib import Path

def docx_to_markdown(docx_file, output_md):
    """Convert a .docx file to a Markdown file and a subfolder of images."""

    folder = str(Path(output_md).parent)
    # 使用输出文件名（不含扩展名）作为图片文件夹名称
    output_filename = Path(output_md).stem
    # 图片文件夹格式：.imgs/文件名/
    image_folder = str(Path(output_md).parent / ".imgs" / output_filename)
    
    doc = docx.Document(docx_file)

    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)
    markdown = []
    image_count = 0
    text = ""

    # save all images
    images = {}
    folder_path = Path(folder)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_info = save_image(rel.target_part, image_folder)
            # 存储相对路径（相对于输出文件夹）和大小信息
            # 使用Path对象计算相对路径
            full_image_path = Path(image_info["path"])
            try:
                relative_path = full_image_path.relative_to(folder_path)
            except ValueError:
                # 如果路径不在folder下，使用原始方式
                relative_path = Path(image_info["path"][len(folder):].lstrip("/\\"))
            images[rel.rId] = {
                "path": str(relative_path).replace("\\", "/"),
                "size": image_info["size"]
            }

    #print("images", images)
    
    for block in doc.element.body:
        if block.tag.endswith('p'):  # Handle paragraphs
            paragraph = paragraphs.pop(0)  # Match current paragraph
            md_paragraph = ""

            style_name = paragraph.style.name

            # 先解析段落内容
            paragraph_content = parse_run(paragraph, images)
            
            # 检查段落是否只包含图片或为空（没有文本内容）
            # 如果内容去除图片标记后没有其他文本，则认为是纯图片段落或空段落
            is_image_only_or_empty = False
            if not paragraph_content.strip():
                # 空段落
                is_image_only_or_empty = True
            else:
                # 移除图片标记后检查是否还有文本
                content_without_images = paragraph_content
                # 移除HTML图片标签
                content_without_images = re.sub(r'<img[^>]*>', '', content_without_images)
                # 移除Markdown图片语法
                content_without_images = re.sub(r'!\[.*?\]\(.*?\)', '', content_without_images)
                # 如果去除图片后没有文本，则认为是纯图片段落
                if not content_without_images.strip():
                    is_image_only_or_empty = True

            #print("Style:", style_name)
            if "List" in style_name:
                # 如果列表项为空或只有图片，不添加列表前缀
                if not is_image_only_or_empty:
                    prefix = get_bullet_point_prefix(paragraph)
                    md_paragraph = prefix  # Markdown syntax for bullet points
            elif "Heading 1" in style_name:
                md_paragraph = "# "
            elif "Heading 2" in style_name:
                md_paragraph = "## "
            elif "Heading 3" in style_name:
                md_paragraph = "### "
            elif "Normal" in style_name:
                md_paragraph = ""
            else:
                print("Unsupported style:", style_name)

            md_paragraph += paragraph_content

            markdown.append(md_paragraph)

        elif block.tag.endswith('tbl'):  # Handle tables (if present)
            table = tables.pop(0)  # Match current table
            table_text = ""
            for i, row in enumerate(table.rows):
                table_text += "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |\n"
                if i == 0:
                    table_text += "| " + " | ".join("---" for _ in row.cells) + " |\n"
                    
            markdown.append(table_text)
            
        #else:
        #    print("Unsupported block:", block)

    # Write to Markdown file
    with open(output_md, "w", encoding="utf-8") as md_file:
        md_file.write("\n\n".join(markdown))


def extract_r_embed(xml_string):
    """
    Extract the value of r:embed from the given XML string.

    :param xml_string: The XML content as a string.
    :return: The value of r:embed or None if not found.
    """
    # Parse the XML
    root = etree.fromstring(xml_string)
    
    # Define the namespaces
    namespaces = {
        'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
        'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        'pic': "http://schemas.openxmlformats.org/drawingml/2006/picture",
    }
    
    # Use XPath to find the <a:blip> element with r:embed
    blip = root.find(".//a:blip", namespaces=namespaces)
    
    # Extract the r:embed attribute value
    if blip is not None:
        return blip.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    return None

def save_image(image_part, output_folder):
    """Save an image to the output folder and return the filename and size."""
    os.makedirs(output_folder, exist_ok=True)
    
    # 获取原始文件扩展名
    original_name = os.path.basename(image_part.partname)
    original_ext = Path(original_name).suffix
    # 如果没有扩展名，尝试从内容推断（简单处理，默认使用 .png）
    if not original_ext:
        original_ext = ".png"
    
    # 使用短唯一标识符：时间戳（6位十六进制）+ 随机数（4位十六进制）
    # 总共10位，比UUID的32位短很多，且足够唯一
    timestamp_hex = format(int(time.time() * 100) & 0xFFFFFF, '06x')  # 6位时间戳（毫秒级）
    random_hex = format(random.randint(0, 0xFFFF), '04x')  # 4位随机数
    unique_filename = f"{timestamp_hex}{random_hex}{original_ext}"
    image_filename = os.path.join(output_folder, unique_filename)
    
    with open(image_filename, "wb") as img_file:
        img_file.write(image_part.blob)
    # 获取图片大小（字节数）
    image_size = len(image_part.blob)
    image_path = str(image_filename).replace("\\", "/")
    return {"path": image_path, "size": image_size}


def get_list_level(paragraph):
    """Determine the level of a bullet point or numbered list item."""
    # Access the raw XML of the paragraph
    p = paragraph._element
    numPr = p.find(".//w:numPr", namespaces=p.nsmap)
    if numPr is not None:
        ilvl = numPr.find(".//w:ilvl", namespaces=p.nsmap)
        if ilvl is not None:
            return int(ilvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"))
    return 0

def get_bullet_point_prefix(paragraph):
    """Determine the Markdown prefix for a bullet point based on its indentation level."""
    level = get_list_level(paragraph)
    return "  " * level + "- "  # Use Markdown syntax for nested lists
    
def parse_run(run, images):
    """Go through document objects recursively and return markdown."""
    sub_parts = list(run.iter_inner_content())
    text = ""
    for s in sub_parts:
        if isinstance(s, str):
            text += s
        elif isinstance(s, docx.text.run.Run):
            text += parse_run(s, images)
        elif isinstance(s, docx.text.hyperlink.Hyperlink):
            text += f"[{s.text}]({s.address})"
        elif isinstance(s, docx.drawing.Drawing):
            rId = extract_r_embed(s._element.xml)
            image_info = images[rId]
            image_path = image_info["path"]
            image_size = image_info["size"]
            # 如果图片小于20KB，使用HTML格式并添加class="icon"
            if image_size < 1024*10:  # 10KB 
                text += f'<img src="./{image_path}" class="img-icon image-with-shadow-base64" />'
            else:
                # 大于等于20KB的图片使用Markdown格式
                text += f"![](./{image_path})"
        else:
            print("unknown run type", s)

    if isinstance(run, docx.text.run.Run):
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        if run.underline:
            text = f"__{text}__"
    return text

